import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import os
from datetime import datetime
import pytz

# File paths
mapping_file = "CFASTR_Category_Mapping_V1.csv"
survey_file = "CFastR_Survey_Data.csv"
thresholds_file = "category_copy_thresholds.csv"
# output_docx will be set dynamically

@st.cache_data
def load_data():
    mapping = pd.read_csv(mapping_file)
    survey = pd.read_csv(survey_file)
    thresholds = pd.read_csv(thresholds_file)
    return mapping, survey, thresholds

def standardize_title(title):
    """Maps variants to canonical job levels for chart grouping."""
    title = str(title).lower()
    if "vp" in title or "vice president" in title or "svp" in title or "evp" in title or "c-suite" in title or "chief" in title:
        return "Executive"
    elif "director" in title:
        return "Director"
    elif "manager" in title:
        return "Manager"
    elif "analyst" in title or "contributor" in title or "associate" in title or "ic" in title or "individual" in title or "staff" in title:
        return "IC"
    else:
        return title.title()

def normalize_key(cat):
    # Normalize for code placeholders
    key = cat.lower().replace(',', '').replace('&', '').replace('-', '').replace('/', '_')
    key = '_'.join(key.split())
    key = key.replace('__', '_')
    return key

def generate_bar_chart(data, title, filename):
    color_list = ['tab:blue', 'tab:orange', 'tab:green', 'tab:red', 
                  'tab:purple', 'tab:brown', 'tab:pink', 'tab:gray', 
                  'tab:olive', 'tab:cyan']
    colors = [color_list[i % len(color_list)] for i in range(len(data))]
    fig, ax = plt.subplots()
    data.plot(kind='bar', ax=ax, color=colors)
    ax.set_title(title)
    plt.tight_layout()
    plt.savefig(filename)
    plt.close(fig)
    return filename

def percent_positive(survey_df, questions, question_polarity):
    positive_count = 0
    total_count = 0
    for q in questions:
        if q not in survey_df.columns:
            continue
        polarity = question_polarity.get(q, 1)  # default to 1 if missing
        responses = pd.to_numeric(survey_df[q], errors='coerce').dropna()
        total_count += responses.shape[0]
        if polarity == 1:
            positive_count += responses.isin([1, 2]).sum()
        else:
            positive_count += responses.isin([4, 5]).sum()
    return (100 * positive_count / total_count) if total_count > 0 else 0

def percent_positive_by_group(survey_df, questions, question_polarity, group_col):
    result = {}
    for group in survey_df[group_col].dropna().unique():
        group_df = survey_df[survey_df[group_col] == group]
        pct = percent_positive(group_df, questions, question_polarity)
        result[group] = pct
    return result

def fill_word_template(doc, data, placeholder_map):
    image_insert_points = []
    for i, para in enumerate(doc.paragraphs):
        for placeholder, key in placeholder_map.items():
            if placeholder in para.text:
                value = data.get(key, "")
                if isinstance(value, str) and value.endswith('.png') and os.path.exists(value):
                    image_insert_points.append((i, value))
                    para.text = para.text.replace(placeholder, "")
                else:
                    para.text = para.text.replace(placeholder, str(value))
    for i, filename in image_insert_points:
        paragraph = doc.paragraphs[i]
        run = paragraph.add_run()
        run.add_picture(filename, width=Inches(5))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for placeholder, key in placeholder_map.items():
                    if placeholder in cell.text:
                        value = data.get(key, "")
                        if isinstance(value, str) and value.endswith('.png') and os.path.exists(value):
                            cell.text = cell.text.replace(placeholder, "")
                            cell.paragraphs[0].add_run().add_picture(value, width=Inches(5))
                        else:
                            cell.text = cell.text.replace(placeholder, str(value))
    return doc

def main():
    st.title("C FASTR Report Generator")

    # --- User Input UI ---
    consultant_name = st.text_input("Consultant Name *")
    company_name = st.text_input("Company Name *")
    client_contact = st.text_input("Client Contact")
    # Use Atlanta timezone for the date
    atlanta_tz = pytz.timezone('America/New_York')
    now_atlanta = datetime.now(atlanta_tz)
    report_date = st.date_input("Report Date", value=now_atlanta.date())
    confidentiality_notice = st.text_area("Confidentiality Notice")
    exec_engagement_summary = st.text_area("Engagement Summary")
    exec_results_summary = st.text_area("Results Summary")
    exec_suggested_actions = st.text_area("Suggested Actions")

    # Validation for required fields
    required_fields = [consultant_name, company_name]
    missing_fields = [f for f, v in zip(["Consultant Name", "Company Name"], required_fields) if not v.strip()]
    if missing_fields:
        st.warning(f"Required fields missing: {', '.join(missing_fields)}")

    if st.button("Generate Report"):
        if missing_fields:
            st.error("Please fill out all required fields before generating the report.")
            st.stop()
        mapping, survey, thresholds = load_data()
        cwd = os.getcwd()
        print(f"[DEBUG] Current working directory: {cwd}")
        print(f"[DEBUG] Columns in survey: {survey.columns.tolist()}")
        print(f"[DEBUG] Columns in mapping: {mapping.columns.tolist()}")

        category_col = "C FASTR Category"
        business_func_col = "Business Function"
        job_level_col = "Title" if "Title" in survey.columns else ("Job Level" if "Job Level" in survey.columns else None)

        # Standardize job level column
        if job_level_col:
            print(f"[DEBUG] Unique values in job_level_col before standardization: {survey[job_level_col].unique()}")
            survey[job_level_col + "_std"] = survey[job_level_col].apply(standardize_title)
            job_level_col_std = job_level_col + "_std"
            job_levels = survey[job_level_col_std].unique()
            print(f"[DEBUG] Unique values in job_level_col after standardization: {job_levels}")
        else:
            job_level_col_std = None
            job_levels = []

        # Build flexible category mappings
        categories = mapping[category_col].unique()
        normalized_to_variants = {}
        for cat in categories:
            norm = normalize_key(cat)
            variants = set([cat, cat.lower(), cat.replace(',', ''), cat.replace(',', ''), norm.replace('_', ' '), norm.replace('_', ', '), norm.replace('_', ',')])
            normalized_to_variants.setdefault(norm, set()).update(variants)
        print(f"[DEBUG] Category normalization mapping:")
        for k, v in normalized_to_variants.items():
            print(f"  {k}: {v}")

        variant_to_normalized = {}
        for norm, variants in normalized_to_variants.items():
            for variant in variants:
                variant_to_normalized[variant.strip().lower()] = norm

        # Polarity mapping
        question_polarity = mapping.set_index("Question Number")["Polarity"].to_dict()

        # Atlanta date for filename
        atlanta_date_str = now_atlanta.strftime("%Y-%m-%d")
        # Safe filename parts
        safe_company = company_name.replace(" ", "_").replace(",", "")
        safe_consultant = consultant_name.replace(" ", "_").replace(",", "")
        output_docx = f"{safe_company}_CultureAssessment_{safe_consultant}_{atlanta_date_str}.docx"

        data = {
            "consultant_name": consultant_name,
            "company_name": company_name,
            "client_contact": client_contact,
            "report_date": report_date.strftime("%Y-%m-%d") if hasattr(report_date, "strftime") else str(report_date),
            "confidentiality_notice": confidentiality_notice,
            "exec_engagement_summary": exec_engagement_summary,
            "exec_results_summary": exec_results_summary,
            "exec_suggested_actions": exec_suggested_actions,
            "conclusion_overview": "Culture is improving.",
            "conclusion_30_60_90": "30/60/90 plan attached.",
            "conclusion_next_30_days": "Focus on feedback.",
            "conclusion_days_31_60": "Accountability training.",
            "conclusion_days_61_90": "Trust-building exercises.",
            "conclusion_metrics_quarterly": "Quarterly survey.",
            "conclusion_closing_thoughts": "Thank you for partnering with us.",
        }

        placeholder_map = {
            "{{ consultant_name }}": "consultant_name",
            "{{ title_company_name }}": "company_name",
            "{{ title_client_contact }}": "client_contact",
            "{{ title_date }}": "report_date",
            "{{ title_confidentiality_notice }}": "confidentiality_notice",
            "{{ exec_engagement_summary }}": "exec_engagement_summary",
            "{{ exec_results_summary }}": "exec_results_summary",
            "{{ exec_suggested_actions }}": "exec_suggested_actions",
            "{{ conclusion_overview }}": "conclusion_overview",
            "{{ conclusion_30_60_90 }}": "conclusion_30_60_90",
            "{{ conclusion_next_30_days }}": "conclusion_next_30_days",
            "{{ conclusion_days_31_60 }}": "conclusion_days_31_60",
            "{{ conclusion_days_61_90 }}": "conclusion_days_61_90",
            "{{ conclusion_metrics_quarterly }}": "conclusion_metrics_quarterly",
            "{{ conclusion_closing_thoughts }}": "conclusion_closing_thoughts",
        }

        # Process every mapping file category
        for i, category in enumerate(categories):
            norm_key = normalize_key(category)
            print(f"[DEBUG] Processing category '{category}' as normalized '{norm_key}'")
            variants = normalized_to_variants[norm_key]
            questions = []
            for v in variants:
                qlist = mapping[mapping[category_col].str.strip().str.lower() == v.strip().lower()]["Question Number"].tolist()
                questions.extend(qlist)
            questions = list(set(questions))  # unique
            print(f"[DEBUG]  Questions found for category '{category}'/'{norm_key}': {questions}")

            if not questions:
                print(f"[DEBUG]    Skipping category '{category}' - no questions found for any variant.")
                continue

            # --- Percent Positive Calculation ---
            overall_percent_positive = percent_positive(survey, questions, question_polarity)
            print(f"[DEBUG]    Overall percent positive for '{category}': {overall_percent_positive:.1f}%")

            # Chart by business function (percent positive)
            by_func = percent_positive_by_group(survey, questions, question_polarity, business_func_col)
            print(f"[DEBUG]    Business Function percent positive data: {by_func}")
            func_bar_filename = f"{norm_key}_bar.png"
            generate_bar_chart(pd.Series(by_func), f"{category} by Business Function (% positive)", func_bar_filename)

            # Chart by job level/title (percent positive)
            if job_level_col_std:
                by_job = percent_positive_by_group(survey, questions, question_polarity, job_level_col_std)
                print(f"[DEBUG]    Job Level percent positive data: {by_job}")
                job_bar_filename = f"{norm_key}_joblevel_bar.png"
                generate_bar_chart(pd.Series(by_job), f"{category} by Job Level (% positive)", job_bar_filename)
            else:
                print(f"[DEBUG]    No job level column found; skipping job level chart for {category}.")
                job_bar_filename = func_bar_filename  # fallback

            # Add placeholders for everything
            placeholder_map[f"{{{{ analysis_{norm_key}_summary }}}}"] = f"{norm_key}_summary"
            placeholder_map[f"{{{{ analysis_{norm_key}_risks }}}}"] = f"{norm_key}_risks"
            placeholder_map[f"{{{{ analysis_{norm_key}_recommendations }}}}"] = f"{norm_key}_recommendations"
            placeholder_map[f"{{{{ {norm_key}_bar }}}}"] = f"{norm_key}_bar_chart"
            placeholder_map[f"{{{{ {norm_key}_joblevel_bar }}}}"] = f"{norm_key}_joblevel_bar_chart"
            placeholder_map[f"{{{{ {norm_key}_pct }}}}"] = f"{norm_key}_pct"

            data[f"{norm_key}_summary"] = f"{overall_percent_positive:.1f}% of responses indicated positive experience in the {category} domain."
            data[f"{norm_key}_risks"] = f"Risks identified in {category}."
            data[f"{norm_key}_recommendations"] = f"Recommend workshops or sessions for {category.lower()}."
            data[f"{norm_key}_bar_chart"] = func_bar_filename
            data[f"{norm_key}_joblevel_bar_chart"] = job_bar_filename
            data[f"{norm_key}_pct"] = f"{overall_percent_positive:.1f}%"

        template_path = "client_report_template.docx"
        doc = Document(template_path)
        doc = fill_word_template(doc, data, placeholder_map)
        doc.save(output_docx)
        st.success(f"Report generated: {output_docx}")

        # Download button for the generated report
        with open(output_docx, "rb") as f:
            docx_bytes = f.read()
        st.download_button(
            label="⬇️ Download Culture Health Report (Word DOCX)",
            data=docx_bytes,
            file_name=output_docx,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

if __name__ == "__main__":
    main()