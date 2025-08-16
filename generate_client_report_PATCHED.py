#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
C FASTR™ client report generator (enhanced)

Inputs (same folder as this script):
  - CFastR_Survey_Data.csv
  - CFASTR_Category_Mapping_V1.csv
  - category_copy_thresholds.csv
  - client_report_template.docx
Optional:
  - consultant_inputs.json  (to override Title/Exec/Conclusion sections)

Outputs:
  - cfastr_charts_by_title/  (PNG charts)
  - Client_CFASTR_Report_Generated.docx

Behavior:
  ✓ Skip "Overall"
  ✓ Normalize categories (ignore commas/dashes; collapse spaces)
  ✓ Top band inclusive of 100% (<= Upper for highest band)
  ✓ Use only mapping-listed question columns
  ✓ Fill Title/Exec/Conclusion placeholders (auto-draft + optional JSON override)
  ✓ Insert per-category Summary/Risks/Recommendations and charts
"""

import os
import re
import json
import unicodedata
import warnings

import pandas as pd
import numpy as np
from pandas.api.types import CategoricalDtype

import matplotlib
matplotlib.use("Agg")  # headless, safe on Mac/servers
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Inches
from textwrap import dedent

# ---------- Config ----------

SURVEY_CSV = "CFastR_Survey_Data.csv"
MAPPING_CSV = "CFASTR_Category_Mapping_V1.csv"
THRESHOLDS_CSV = "category_copy_thresholds.csv"
TEMPLATE_DOCX = "client_report_template.docx"
OUTPUT_DOCX = "Client_CFASTR_Report_Generated.docx"

TITLE_COL = "Title"
FUNCTION_COL = "Business Function"

COPY_CATEGORY_COL = "C FASTR Category"
COPY_LOWER_THRESHOLD_COL = "Lower Threshold Value"
COPY_UPPER_THRESHOLD_COL = "Upper Threshold Value"
COPY_SUMMARY_COL = "Copy for Summary Section"
COPY_RISKS_COL = "Copy for Risks Section"
COPY_RECOMMENDATIONS_COL = "Copy for Recommendations Section"

CHARTS_DIR = "cfastr_charts_by_title"

PREFERRED_TITLE_ORDER = [
    "Chief People Officer",
    "Head of People and Culture",
    "VP",
    "Director",
    "Manager",
    "IC",
]

# ---------- Helpers ----------

def normalize_category_key(s: str) -> str:
    """Normalize category for matching: ignore commas/dashes; collapse spaces; lowercase."""
    if s is None:
        return ""
    t = unicodedata.normalize("NFKC", str(s)).strip()
    t = t.replace("—", "-").replace("–", "-")
    t = re.sub(r"[,-]", " ", t)           # commas & dashes -> space
    t = re.sub(r"\s+", " ", t)
    return t.lower()

def category_to_token(s: str) -> str:
    """Convert category to token for placeholders/chart filenames (A–Z0–9 + underscores)."""
    if s is None:
        return ""
    t = unicodedata.normalize("NFKC", str(s)).strip().upper()
    t = t.replace("—", " ").replace("–", " ").replace("-", " ").replace(",", " ")
    t = re.sub(r"[^A-Z0-9]+", "_", t).strip("_")
    return t

def ensure_dir(path: str):
    if not os.path.isdir(path):
        os.makedirs(path, exist_ok=True)

def read_csv_required(path: str, label: str) -> pd.DataFrame:
    if not os.path.exists(path):
        raise FileNotFoundError(f"Required file not found: {path} ({label})")
    df = pd.read_csv(path)
    # Strip BOM and whitespace from headers to avoid KeyErrors
    df.columns = [c.replace("\ufeff", "").strip() for c in df.columns]
    return df

def positive_mask(series_like5: pd.Series) -> pd.Series:
    """Positive if score >= 4."""
    return series_like5 >= 4

def apply_polarity(values: pd.Series, polarity: int) -> pd.Series:
    """If polarity == -1, invert 1..5 scale to 5..1."""
    if polarity == -1:
        return 6 - values
    return values

def get_category_copy(category_name: str, pct_positive: float, thresholds_df: pd.DataFrame) -> dict:
    """
    Select copy row where:
      - Lower <= pct < Upper for intermediate bands
      - For TOP band (max Upper): Lower <= pct <= Upper  (inclusive of 100%)
    """
    if thresholds_df.empty:
        return {"summary": "", "risks": "", "recommendations": ""}

    df = thresholds_df.copy()
    df.columns = [c.replace("\ufeff", "").strip() for c in df.columns]
    df["_norm_cat"] = df[COPY_CATEGORY_COL].map(normalize_category_key)
    norm_cat = normalize_category_key(category_name)
    cat_df = df[df["_norm_cat"] == norm_cat].copy()

    if cat_df.empty:
        warnings.warn(f"No thresholds for category: {category_name}")
        return {"summary": "", "risks": "", "recommendations": ""}

    # numeric
    cat_df[COPY_LOWER_THRESHOLD_COL] = pd.to_numeric(cat_df[COPY_LOWER_THRESHOLD_COL], errors="coerce")
    cat_df[COPY_UPPER_THRESHOLD_COL] = pd.to_numeric(cat_df[COPY_UPPER_THRESHOLD_COL], errors="coerce")
    cat_df = cat_df.dropna(subset=[COPY_LOWER_THRESHOLD_COL, COPY_UPPER_THRESHOLD_COL])
    if cat_df.empty:
        return {"summary": "", "risks": "", "recommendations": ""}

    top_upper = cat_df[COPY_UPPER_THRESHOLD_COL].max()

    inter_mask = (
        (cat_df[COPY_LOWER_THRESHOLD_COL] <= pct_positive) &
        (pct_positive < cat_df[COPY_UPPER_THRESHOLD_COL]) &
        (cat_df[COPY_UPPER_THRESHOLD_COL] < top_upper)
    )
    top_mask = (
        (cat_df[COPY_UPPER_THRESHOLD_COL] == top_upper) &
        (cat_df[COPY_LOWER_THRESHOLD_COL] <= pct_positive) &
        (pct_positive <= cat_df[COPY_UPPER_THRESHOLD_COL])
    )

    match_df = pd.concat([cat_df[inter_mask], cat_df[top_mask]], axis=0)
    if match_df.empty:
        warnings.warn(f"No threshold matched for {category_name} at {pct_positive:.1f}%")
        return {"summary": "", "risks": "", "recommendations": ""}

    row = match_df.sort_values(by=COPY_LOWER_THRESHOLD_COL, ascending=False).iloc[0]

    return {
        "summary": row.get(COPY_SUMMARY_COL, "") or "",
        "risks": row.get(COPY_RISKS_COL, "") or "",
        "recommendations": row.get(COPY_RECOMMENDATIONS_COL, "") or "",
    }

def replace_placeholder_text(paragraph, placeholder: str, text: str):
    """Replace a placeholder token in a paragraph with plain text (keeps paragraph style)."""
    if placeholder not in paragraph.text:
        return
    # Rebuild paragraph text cleanly (python-docx run-safe replace)
    inline = paragraph.runs
    new_text = paragraph.text.replace(placeholder, text)
    for i in range(len(inline) - 1, -1, -1):
        r = inline[i]._r
        paragraph._p.remove(r)
    paragraph.add_run(new_text)

def insert_chart_at_placeholder(document: Document, placeholder: str, image_path: str, width_in: float = 6.0):
    """Replace a placeholder with an image (search paragraphs and table cells)."""
    # paragraphs
    for paragraph in document.paragraphs:
        if placeholder in paragraph.text:
            replace_placeholder_text(paragraph, placeholder, "")
            run = paragraph.add_run()
            run.add_picture(image_path, width=Inches(width_in))
    # tables
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if placeholder in paragraph.text:
                        replace_placeholder_text(paragraph, placeholder, "")
                        run = paragraph.add_run()
                        run.add_picture(image_path, width=Inches(width_in))

def load_consultant_inputs(path="consultant_inputs.json"):
    """Return dict or None if no JSON present."""
    if not os.path.exists(path):
        return None
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)

def replace_token_everywhere(doc: Document, token: str, text: str):
    """Replace token in all paragraphs and tables."""
    # paragraphs
    for p in doc.paragraphs:
        if token in p.text:
            replace_placeholder_text(p, token, text)
    # tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if token in p.text:
                        replace_placeholder_text(p, token, text)

def _fmt_pct(x):
    return f"{x:.0f}%"

def _fmt_score(x):
    return f"{x:.2f}"

def _list3(rows: pd.DataFrame, name_col: str, val_col: str):
    """Format top/bottom 3 like: A (75%), B (72%), C (69%)."""
    out = []
    for _, r in rows.iterrows():
        nm = str(r[name_col])
        val = r[val_col]
        if pd.isna(val):
            continue
        out.append(f"{nm} ({_fmt_pct(val)})")
    return ", ".join(out)

def suggest_conclusion_text(by_cat: pd.DataFrame,
                            avg_by_title: pd.DataFrame,
                            avg_by_func: pd.DataFrame):
    """
    Returns:
      overview, plan_30_60_90, next_30, days_31_60, days_61_90, metrics_quarterly, closing_thoughts
    """
    # Overall indicators
    total_n = int(by_cat["N"].sum())
    overall_pos = 100.0 * by_cat["PosN"].sum() / max(1, by_cat["N"].sum())
    overall_avg = (by_cat["Avg"] * by_cat["N"]).sum() / max(1, by_cat["N"].sum())

    # Top/bottom by % positive (relax N threshold if needed)
    bc = by_cat.copy()
    bc = bc[bc["N"] >= 1]
    strengths = bc.sort_values("Percentage Positive", ascending=False).head(3)
    focus = bc.sort_values("Percentage Positive", ascending=True).head(3)

    # Dispersion by level
    gap_title = []
    for cat, g in avg_by_title.groupby(COPY_CATEGORY_COL):
        vals = g["AdjScore"].dropna().tolist()
        if len(vals) >= 2:
            gap = max(vals) - min(vals)
            gap_title.append((cat, gap))
    gap_title = sorted(gap_title, key=lambda t: t[1], reverse=True)[:2]

    # Dispersion by function
    gap_func = []
    for cat, g in avg_by_func.groupby(COPY_CATEGORY_COL):
        vals = g["AdjScore"].dropna().tolist()
        if len(vals) >= 2:
            gap = max(vals) - min(vals)
            gap_func.append((cat, gap))
    gap_func = sorted(gap_func, key=lambda t: t[1], reverse=True)[:2]

    # Build Overview
    overview = dedent(f"""
    This report summarizes organizational signals across C FASTR categories based on {total_n} scored responses.
    Overall, {_fmt_pct(overall_pos)} of responses were positive (average score {_fmt_score(overall_avg)} out of 5).

    Strengths observed: {_list3(strengths, COPY_CATEGORY_COL, "Percentage Positive")}.

    Priority focus areas: {_list3(focus, COPY_CATEGORY_COL, "Percentage Positive")}.

    Notable variation by level and function:
    - By level (Title): {", ".join([f"{c} (gap {_fmt_score(g)})" for c,g in gap_title]) or "no pronounced gaps"}.
    - By function: {", ".join([f"{c} (gap {_fmt_score(g)})" for c,g in gap_func]) or "no pronounced gaps"}.
    """).strip()

    # Action hints keyed by normalized category
    ACTION_HINTS = {
        "feedback receiving": {
            "30": "Establish a simple feedback ritual (e.g., ‘start/stop/continue’) at the team level; model receivership skills in staff meetings.",
            "60": "Train managers on coaching micro-skills; implement a monthly ‘feedback Friday’ cadence.",
            "90": "Embed feedback into 1:1 templates and quarterly reviews; recognize teams demonstrating high-quality feedback loops."
        },
        "feedback giving": {
            "30": "Introduce a shared definition of ‘good feedback’ and a 2x2 example library (behavior, impact, request).",
            "60": "Run manager practice labs; add peer feedback into sprint retros or project post-mortems.",
            "90": "Integrate feedback expectations into performance criteria; publish anonymized ‘great feedback’ exemplars."
        },
        "accountability": {
            "30": "Clarify owners and due dates on top cross-functional deliverables; publish a visible RAID log.",
            "60": "Adopt a weekly commitments review (what was promised/what shipped); normalize ask-by/commit-by language.",
            "90": "Connect commitments to OKRs; review slippage trends and remove structural blockers."
        },
        "sensitivity": {
            "30": "Set team norms on respectful debate and meeting hygiene; introduce ‘one-mic’ and ‘assume positive intent’.",
            "60": "Offer manager toolkits for navigating hard conversations; audit meeting load & inclusion signals.",
            "90": "Reinforce norms via recognition; ensure retro actions close the loop on psychological safety issues."
        },
        "relationship focus": {
            "30": "Create lightweight cross-team touchpoints (buddy program or coffee chats); add intros in key workflows.",
            "60": "Stand up a quarterly cross-functional forum to unblock work; measure collaboration sentiment.",
            "90": "Formalize partnership charters where dependencies are chronic; track cycle-time improvements."
        },
        "trust": {
            "30": "Leaders share context and constraints in plain language; schedule listening sessions with Q&A.",
            "60": "Publish decision rationales post-launch; add ‘what we learned/changed’ sections to reviews.",
            "90": "Commit to a predictable comms cadence; assess trust drivers quarterly and adjust rituals."
        },
        "collusion": {
            "30": "Name and neutralize ‘good soldier’ behaviors; encourage dissent using red-team reviews.",
            "60": "Instrument escalation paths and celebrate principled pushback.",
            "90": "Bake in pre-mortems and decision logs to keep healthy tension visible."
        },
    }

    # Choose anchor categories from bottom-of-list and largest level gap
    norm = lambda s: normalize_category_key(s)
    focus_norms = [norm(c) for c in focus[COPY_CATEGORY_COL].tolist()]
    gap_anchor = norm(gap_title[0][0]) if gap_title else None

    anchors = []
    for c in focus_norms:
        if c not in anchors:
            anchors.append(c)
    if gap_anchor and gap_anchor not in anchors:
        anchors.append(gap_anchor)
    anchors = anchors[:3]  # keep concise

    def plan_line(cat_norm, day):
        hints = ACTION_HINTS.get(cat_norm) or {}
        return hints.get(day, "Stand up a light-weight inspect-and-adapt loop; define owner, cadence, and success signal.")

    # Build 30/60/90 lines grounded in anchors
    lines_30 = [f"• {plan_line(a, '30')}" for a in anchors] or ["• Socialize findings; align on 2–3 priority experiments."]
    lines_60 = [f"• {plan_line(a, '60')}" for a in anchors] or ["• Train managers on targeted rituals; establish measurement baselines."]
    lines_90 = [f"• {plan_line(a, '90')}" for a in anchors] or ["• Integrate successful rituals into operating cadence; report deltas quarterly."]

    plan_3090 = dedent(f"""
    30 days — Quick wins and clarity
    {chr(10).join(lines_30)}

    60 days — Build durable habits
    {chr(10).join(lines_60)}

    90 days — Scale and measure
    {chr(10).join(lines_90)}
    """).strip()

    # Split-out sections (for dedicated placeholders)
    next_30     = "\n".join(lines_30)
    days_31_60  = "\n".join(lines_60)
    days_61_90  = "\n".join(lines_90)

    metrics_quarterly = (
        "• Overall % positive and per-category deltas\n"
        "• Gap by Title and by Business Function (max–min average)\n"
        "• Feedback cadence (frequency/coverage), 1:1 completion, retro/action closure\n"
        "• Time-to-decision / time-to-unblock on cross-functional work\n"
        "• Participation/response rates; trust/safety micro-indices"
    )

    closing = (
        "We’re seeing clear strengths to amplify and a focused set of habits to tighten. "
        "With visible ownership, simple rituals, and steady measurement, momentum typically shows up inside the first quarter. "
        "Maintain a 30-day inspect-and-adapt loop and a quarterly read-out on the metrics above to keep progress durable."
    )

    return overview, plan_3090, next_30, days_31_60, days_61_90, metrics_quarterly, closing

# ---------- Main ----------

def main():
    # Load inputs
    survey = read_csv_required(SURVEY_CSV, "Survey data")
    mapping = read_csv_required(MAPPING_CSV, "Question -> Category mapping")
    thresholds = read_csv_required(THRESHOLDS_CSV, "Copy thresholds")

    # Basic column checks
    for col in (TITLE_COL, FUNCTION_COL):
        if col not in survey.columns:
            raise KeyError(f"Survey CSV must include '{col}' column.")

    if "Question Number" not in mapping.columns or COPY_CATEGORY_COL not in mapping.columns:
        raise KeyError("Mapping CSV must include 'Question Number' and 'C FASTR Category'.")

    # Polarity default = 1
    if "Polarity" not in mapping.columns:
        mapping["Polarity"] = 1
    mapping["Polarity"] = pd.to_numeric(mapping["Polarity"], errors="coerce").fillna(1).astype(int)

    # Internal normalized category (for filtering/skip)
    mapping["_norm_cat"] = mapping[COPY_CATEGORY_COL].map(normalize_category_key)

    # Use ONLY mapped questions (ignore extras in survey)
    mapped_questions = [q for q in mapping["Question Number"].astype(str).tolist() if q in survey.columns]
    if not mapped_questions:
        raise ValueError("No question columns from Mapping were found in Survey CSV.")

    # Melt survey for mapped questions
    df_melt = survey[[TITLE_COL, FUNCTION_COL] + mapped_questions].melt(
        id_vars=[TITLE_COL, FUNCTION_COL],
        value_vars=mapped_questions,
        var_name="Question Number",
        value_name="Score"
    )
    df_melt["Score"] = pd.to_numeric(df_melt["Score"], errors="coerce")
    df_melt = df_melt.dropna(subset=["Score"])

    # Join mapping (to get category + polarity), then skip 'Overall'
    df = df_melt.merge(
        mapping[["Question Number", COPY_CATEGORY_COL, "_norm_cat", "Polarity"]],
        on="Question Number",
        how="left"
    )
    df = df[df["_norm_cat"] != "overall"]

    # Apply polarity to scores
    df["AdjScore"] = df.apply(lambda r: apply_polarity(r["Score"], r["Polarity"]), axis=1)

    # ---------- Compute metrics ----------

    pos = positive_mask(df["AdjScore"])

    by_cat = (
        df.assign(Pos=pos)
          .groupby(COPY_CATEGORY_COL)
          .agg(
              N=("AdjScore", "count"),
              PosN=("Pos", "sum"),
              Avg=("AdjScore", "mean"),
          )
          .reset_index()
    )
    by_cat["Percentage Positive"] = np.where(
        by_cat["N"] > 0,
        100.0 * by_cat["PosN"] / by_cat["N"],
        np.nan
    )

    avg_by_func = (
        df.groupby([COPY_CATEGORY_COL, FUNCTION_COL])["AdjScore"]
          .mean()
          .reset_index()
    )
    avg_by_title = (
        df.groupby([COPY_CATEGORY_COL, TITLE_COL])["AdjScore"]
          .mean()
          .reset_index()
    )

    # ----- Auto-suggest Conclusion text from data -----
    (auto_overview,
     auto_3090,
     auto_next_30,
     auto_days_31_60,
     auto_days_61_90,
     auto_metrics_qtr,
     auto_closing) = suggest_conclusion_text(by_cat, avg_by_title, avg_by_func)

    # ---------- Resolve copy for each category ----------
    resolved_copy = {}
    for _, row in by_cat.iterrows():
        category_name = row[COPY_CATEGORY_COL]
        pct = float(row["Percentage Positive"]) if pd.notna(row["Percentage Positive"]) else 0.0
        resolved_copy[category_name] = get_category_copy(category_name, pct, thresholds)

    # ---------- Generate charts ----------
    ensure_dir(CHARTS_DIR)

    actual_titles = avg_by_title[TITLE_COL].dropna().unique().tolist()
    title_order = [t for t in PREFERRED_TITLE_ORDER if t in actual_titles] or sorted(actual_titles)

    unique_categories = [
        c for c in by_cat[COPY_CATEGORY_COL].unique().tolist()
        if normalize_category_key(c) != "overall"
    ]

    for category in unique_categories:
        token = category_to_token(category)

        # Chart 1: by Business Function
        sub_func = avg_by_func[avg_by_func[COPY_CATEGORY_COL] == category].copy()
        if not sub_func.empty:
            sub_func = sub_func.sort_values(by="AdjScore", ascending=True)
            plt.figure(figsize=(8, max(2.5, 0.35 * len(sub_func))))
            plt.barh(sub_func[FUNCTION_COL], sub_func["AdjScore"])
            plt.xlabel("Average Score (1–5)")
            plt.title(f"{category}: Avg Score by Business Function")
            fname1 = os.path.join(CHARTS_DIR, f"avg_score_{token.lower()}_by_business_function.png")
            plt.tight_layout()
            plt.savefig(fname1, dpi=200)
            plt.close()

        # Chart 2: by Title (Job Level)
        sub_title = avg_by_title[avg_by_title[COPY_CATEGORY_COL] == category].copy()
        if not sub_title.empty:
            cat_type = CategoricalDtype(categories=title_order, ordered=True)
            sub_title[TITLE_COL] = sub_title[TITLE_COL].astype(cat_type)
            sub_title = sub_title.sort_values(by=TITLE_COL)
            plt.figure(figsize=(8, max(2.5, 0.35 * len(sub_title))))
            plt.barh(sub_title[TITLE_COL].astype(str), sub_title["AdjScore"])
            plt.xlabel("Average Score (1–5)")
            plt.title(f"{category}: Avg Score by Level")
            fname2 = os.path.join(CHARTS_DIR, f"avg_score_{token.lower()}_by_job_level.png")
            plt.tight_layout()
            plt.savefig(fname2, dpi=200)
            plt.close()

    # ---------- Compose Word document ----------
    if not os.path.exists(TEMPLATE_DOCX):
        raise FileNotFoundError(f"Template not found: {TEMPLATE_DOCX}")

    doc = Document(TEMPLATE_DOCX)

    # Top sections: Title/Exec/Conclusion (JSON overrides auto-draft)
    ci = load_consultant_inputs("consultant_inputs.json")

    title_company = title_contact = title_date = title_notice = ""
    exec_eng = exec_res = exec_act = ""

    # Defaults to auto-draft for Conclusion
    conc_overview    = auto_overview
    conc_3090        = auto_3090
    conc_next_30     = auto_next_30
    conc_days_31_60  = auto_days_31_60
    conc_days_61_90  = auto_days_61_90
    conc_metrics_qtr = auto_metrics_qtr
    conc_closing     = auto_closing

    if ci:
        tp = ci.get("title_page", {})
        ex = ci.get("exec_summary", {})
        co = ci.get("conclusion", {})

        title_company = (tp.get("company_name") or "").strip()
        title_contact = (tp.get("client_contact") or "").strip()
        title_date    = (tp.get("report_date") or "").strip()
        title_notice  = (tp.get("confidentiality_notice") or "").strip()
        if not title_notice:
            cn = title_company if title_company else "[Company name]"
            title_notice = (f"This report contains confidential and proprietary information intended solely for {cn}. "
                            f"Do not distribute, reproduce, or share without written permission from The Transformation Guild and {cn}.")

        exec_eng = (ex.get("engagement_summary") or "").strip()
        exec_res = (ex.get("results_summary") or "").strip()
        exec_act = (ex.get("suggested_actions") or "").strip()

        # Consultant overrides for Conclusion (optional)
        if (co.get("overview") or "").strip():
            conc_overview = co["overview"].strip()
        if (co.get("thirty_sixty_ninety") or "").strip():
            conc_3090 = co["thirty_sixty_ninety"].strip()
        if (co.get("next_30_days") or "").strip():
            conc_next_30 = co["next_30_days"].strip()
        if (co.get("days_31_60") or "").strip():
            conc_days_31_60 = co["days_31_60"].strip()
        if (co.get("days_61_90") or "").strip():
            conc_days_61_90 = co["days_61_90"].strip()
        if (co.get("metrics_quarterly") or "").strip():
            conc_metrics_qtr = co["metrics_quarterly"].strip()
        if (co.get("closing_thoughts") or "").strip():
            conc_closing = co["closing_thoughts"].strip()

    tokens = {
        # Title page
        "[TITLE_COMPANY_NAME]": title_company,
        "[TITLE_CLIENT_CONTACT]": title_contact,
        "[TITLE_DATE]": title_date,
        "[TITLE_CONFIDENTIALITY_NOTICE]": title_notice,
        # Executive summary
        "[EXEC_ENGAGEMENT_SUMMARY]": exec_eng,
        "[EXEC_RESULTS_SUMMARY]":    exec_res,
        "[EXEC_SUGGESTED_ACTIONS]":  exec_act,
        # Conclusion (existing)
        "[CONCLUSION_OVERVIEW]":     conc_overview,
        "[CONCLUSION_30_60_90]":     conc_3090,
        # Conclusion (new split sections)
        "[CONCLUSION_NEXT_30_DAYS]":      conc_next_30,
        "[CONCLUSION_DAYS_31_60]":        conc_days_31_60,
        "[CONCLUSION_DAYS_61_90]":        conc_days_61_90,
        "[CONCLUSION_METRICS_QUARTERLY]": conc_metrics_qtr,
        "[CONCLUSION_CLOSING_THOUGHTS]":  conc_closing,
    }
    for tok, val in tokens.items():
        replace_token_everywhere(doc, tok, val)

    # Per-category analysis + charts
    for category in unique_categories:
        token = category_to_token(category)
        copy_block = resolved_copy.get(category, {}) or {}

        # Text placeholders
        summary_ph = f"[ANALYSIS_{token}_SUMMARY]"
        risks_ph   = f"[ANALYSIS_{token}_RISKS]"
        recs_ph    = f"[ANALYSIS_{token}_RECOMMENDATIONS]"

        # Replace in paragraphs and tables
        replace_token_everywhere(doc, summary_ph, copy_block.get("summary", ""))
        replace_token_everywhere(doc, risks_ph,   copy_block.get("risks", ""))
        replace_token_everywhere(doc, recs_ph,    copy_block.get("recommendations", ""))

        # Chart placeholders
        chart1_ph = f"[CHART_{token}_BY_BUSINESS_FUNCTION]"
        chart2_ph = f"[CHART_{token}_BY_LEVEL]"
        chart1_fn = os.path.join(CHARTS_DIR, f"avg_score_{token.lower()}_by_business_function.png")
        chart2_fn = os.path.join(CHARTS_DIR, f"avg_score_{token.lower()}_by_job_level.png")

        if os.path.exists(chart1_fn):
            insert_chart_at_placeholder(doc, chart1_ph, chart1_fn, width_in=6.0)
        if os.path.exists(chart2_fn):
            insert_chart_at_placeholder(doc, chart2_ph, chart2_fn, width_in=6.0)

    doc.save(OUTPUT_DOCX)
    print(f"Report written: {OUTPUT_DOCX}")

if __name__ == "__main__":
    main()
